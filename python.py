import tkinter as tk
from tkinter import filedialog, messagebox
import os
from docx import Document

class ContractGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Contract Generator")

        # Dossiers des templates et des contrats générés
        self.CONTRACT_TEMPLATES_FOLDER = "./contrat"
        self.NEW_CONTRACTS_FOLDER = "./nouveau_contrat"

        # Vérifier si le dossier des templates existe
        if not os.path.exists(self.CONTRACT_TEMPLATES_FOLDER):
            os.makedirs(self.CONTRACT_TEMPLATES_FOLDER)

        # Vérifier si le dossier des nouveaux contrats existe
        if not os.path.exists(self.NEW_CONTRACTS_FOLDER):
            os.makedirs(self.NEW_CONTRACTS_FOLDER)

        # Initialize variables
        self.lawyer_name_var = tk.StringVar()
        self.lawyer_phone_var = tk.StringVar()
        self.lawyer_email_var = tk.StringVar()
        self.lawyer_address_var = tk.StringVar()
        self.assistant_name_var = tk.StringVar()
        self.assistant_phone_var = tk.StringVar()
        self.file_number_var = tk.StringVar()
        self.matter_description_var = tk.StringVar()
        self.client_name_var = tk.StringVar()
        self.client_address_var = tk.StringVar()
        self.opposing_party_name_var = tk.StringVar()
        self.opposing_party_phone_var = tk.StringVar()
        self.court_file_number_var = tk.StringVar()
        self.court_registry_location_var = tk.StringVar()
        self.contract_type_var = tk.StringVar()

        # Create UI elements
        tk.Label(root, text="Lawyer Name").grid(row=0)
        tk.Entry(root, textvariable=self.lawyer_name_var).grid(row=0, column=1)

        tk.Label(root, text="Lawyer Phone").grid(row=1)
        tk.Entry(root, textvariable=self.lawyer_phone_var).grid(row=1, column=1)

        tk.Label(root, text="Lawyer Email").grid(row=2)
        tk.Entry(root, textvariable=self.lawyer_email_var).grid(row=2, column=1)

        tk.Label(root, text="Lawyer Address").grid(row=3)
        tk.Entry(root, textvariable=self.lawyer_address_var).grid(row=3, column=1)

        tk.Label(root, text="Assistant Name").grid(row=4)
        tk.Entry(root, textvariable=self.assistant_name_var).grid(row=4, column=1)

        tk.Label(root, text="Assistant Phone").grid(row=5)
        tk.Entry(root, textvariable=self.assistant_phone_var).grid(row=5, column=1)

        tk.Label(root, text="File Number").grid(row=6)
        tk.Entry(root, textvariable=self.file_number_var).grid(row=6, column=1)

        tk.Label(root, text="Matter Description").grid(row=7)
        tk.Entry(root, textvariable=self.matter_description_var).grid(row=7, column=1)

        tk.Label(root, text="Client Name").grid(row=8)
        tk.Entry(root, textvariable=self.client_name_var).grid(row=8, column=1)

        tk.Label(root, text="Client Address").grid(row=9)
        tk.Entry(root, textvariable=self.client_address_var).grid(row=9, column=1)

        tk.Label(root, text="Opposing Party Name").grid(row=10)
        tk.Entry(root, textvariable=self.opposing_party_name_var).grid(row=10, column=1)

        tk.Label(root, text="Opposing Party Phone").grid(row=11)
        tk.Entry(root, textvariable=self.opposing_party_phone_var).grid(row=11, column=1)

        tk.Label(root, text="Court File Number").grid(row=12)
        tk.Entry(root, textvariable=self.court_file_number_var).grid(row=12, column=1)

        tk.Label(root, text="Court Registry Location").grid(row=13)
        tk.Entry(root, textvariable=self.court_registry_location_var).grid(row=13, column=1)

        tk.Label(root, text="Type of Contract").grid(row=14)
        self.contract_type_var.set("Select contract type")  # Set initial value
        self.contract_type_menu = tk.OptionMenu(root, self.contract_type_var, *self.get_contract_types())
        self.contract_type_menu.grid(row=14, column=1)

        tk.Button(root, text='Generate Contract', command=self.generate_contract).grid(row=15, column=0, pady=4)
        tk.Button(root, text='Preview Contract', command=self.preview_contract).grid(row=15, column=1, pady=4)
        tk.Button(root, text='Quit', command=root.quit).grid(row=15, column=2, pady=4)

    def get_contract_types(self):
        # Liste des types de contrats basée sur les fichiers dans le dossier des templates
        try:
            contract_types = [os.path.splitext(filename)[0] for filename in os.listdir(self.CONTRACT_TEMPLATES_FOLDER) if filename.endswith('.docx')]
        except Exception as e:
            messagebox.showerror("Error", f"Failed to list contract templates: {e}")
            contract_types = []
        return contract_types

    def get_mappings(self):
        return {
            "{LawyerName}": self.lawyer_name_var.get(),
            "{LawyerPhone}": self.lawyer_phone_var.get(),
            "{LawyerEmail}": self.lawyer_email_var.get(),
            "{LawyerStreetAddress}": self.lawyer_address_var.get(),
            "{AssistantName}": self.assistant_name_var.get(),
            "{AssistantPhone}": self.assistant_phone_var.get(),
            "{FileNumber}": self.file_number_var.get(),
            "{MatterDescription}": self.matter_description_var.get(),
            "{ClientName}": self.client_name_var.get(),
            "{ClientAddress}": self.client_address_var.get(),
            "{OpposingPartyName}": self.opposing_party_name_var.get(),
            "{OpposingPartyPhone}": self.opposing_party_phone_var.get(),
            "{CourtFileNumber}": self.court_file_number_var.get(),
            "{CourtRegistryLocation}": self.court_registry_location_var.get(),
            "{TypeofContract}": self.contract_type_var.get()
        }

    def generate_contract(self):
        # Mapping for placeholder replacements
        mapping = self.get_mappings()

        # Select template based on contract type
        template_filename = filedialog.askopenfilename(
            initialdir=self.CONTRACT_TEMPLATES_FOLDER,
            title="Select template",
            filetypes=(("Word files", "*.docx"), ("all files", "*.*"))
        )
        if not template_filename:
            messagebox.showerror("Error", "No template selected")
            return

        # Load the document template
        doc = Document(template_filename)

        # Replace placeholders with actual values
        for p in doc.paragraphs:
            for key, value in mapping.items():
                if key in p.text:
                    p.text = p.text.replace(key, value)

        # Save the filled contract in a new file
        output_dir = self.NEW_CONTRACTS_FOLDER
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"{self.file_number_var.get()}_{self.client_name_var.get()}.docx")
        doc.save(output_path)

        messagebox.showinfo("Success", f"Contract generated and saved at {output_path}")

    def preview_contract(self):
        # Mapping for placeholder replacements
        mapping = self.get_mappings()

        # Select template based on contract type
        template_filename = filedialog.askopenfilename(
            initialdir=self.CONTRACT_TEMPLATES_FOLDER,
            title="Select template",
            filetypes=(("Word files", "*.docx"), ("all files", "*.*"))
        )
        if not template_filename:
            messagebox.showerror("Error", "No template selected")
            return

        # Load the document template
        doc = Document(template_filename)

        # Replace placeholders with actual values
        for p in doc.paragraphs:
            for key, value in mapping.items():
                if key in p.text:
                    p.text = p.text.replace(key, value)

        # Save the filled contract in a new file for preview
        preview_dir = self.NEW_CONTRACTS_FOLDER
        os.makedirs(preview_dir, exist_ok=True)
        preview_path = os.path.join(preview_dir, f"{self.file_number_var.get()}_{self.client_name_var.get()}_preview.docx")
        doc.save(preview_path)

        messagebox.showinfo("Success", f"Preview contract generated and saved at {preview_path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = ContractGeneratorApp(root)
    root.mainloop()
